import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

# ─────────────────────────────────────────
#  HELPER FUNCTIONS
# ─────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margin(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for tag, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        n = OxmlElement(tag)
        n.set(qn('w:w'), str(val))
        n.set(qn('w:type'), 'dxa')
        tcMar.append(n)
    tcPr.append(tcMar)

def add_h1(doc, text, color=(0x00, 0x33, 0x66)):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    return p

def add_h2(doc, text, color=(0x00, 0x55, 0x99)):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    return p

def add_h3(doc, text):
    p = doc.add_heading(level=3)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_bullet(doc, label, text, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
    p.paragraph_format.space_after = Pt(4)
    if label:
        r1 = p.add_run(f"{label}: ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
    r2 = p.add_run(text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    return p

def add_note_box(doc, text, color_hex="FFF3CD", border_color=(0xFF, 0xA0, 0x00)):
    """Add a shaded note/callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    pPr.append(shading)
    r = p.add_run(f"💡  {text}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x22, 0x00)
    return p

def add_table(doc, headers, rows, col_widths=None, header_color="003366"):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = 'Table Grid'
    # Header row
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_color)
        set_cell_margin(hdr[i])
        if col_widths:
            hdr[i].width = col_widths[i]
        p = hdr[i].paragraphs[0]
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    # Data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            set_cell_margin(cells[ci])
            if col_widths:
                cells[ci].width = col_widths[ci]
            run = cells[ci].paragraphs[0].runs[0]
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            if ri % 2 == 1:
                set_cell_bg(cells[ci], "F0F4FA")
    return t

def page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────
#  BUILD DOCUMENT
# ─────────────────────────────────────────
def build():
    doc = docx.Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    # Base style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ══════════════════════════════════════════════
    #  TRANG BÌA
    # ══════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    t_run = title_p.add_run("TÀI LIỆU LÝ THUYẾT DỰ ÁN")
    t_run.font.name = 'Times New Roman'
    t_run.font.size = Pt(20)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    sub1_p = doc.add_paragraph()
    sub1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub1_run = sub1_p.add_run("PHÁT HIỆN ACCESS POINT GIẢ MẠO (ROGUE AP)\nVÀ TẤN CÔNG EVIL TWIN\nBẰNG HỌC MÁY CÓ GIÁM SÁT")
    sub1_run.font.name = 'Times New Roman'
    sub1_run.font.size = Pt(16)
    sub1_run.font.bold = True
    sub1_run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)

    doc.add_paragraph("\n" * 3)

    intro_p = doc.add_paragraph()
    intro_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_run = intro_p.add_run(
        "Tài liệu này dành cho sinh viên — giải thích toàn bộ các khái niệm lý thuyết,\n"
        "công nghệ, thuật toán và quy trình nghiên cứu của dự án một cách dễ hiểu nhất."
    )
    intro_run.font.name = 'Times New Roman'
    intro_run.font.size = Pt(12)
    intro_run.font.italic = True
    intro_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    page_break(doc)

    # ══════════════════════════════════════════════
    #  MỤC LỤC (thủ công)
    # ══════════════════════════════════════════════
    add_h1(doc, "MỤC LỤC")
    toc_items = [
        ("Chương 1", "Mạng Wi-Fi là gì? Hoạt động ra sao?"),
        ("Chương 2", "Access Point (AP) và các loại AP"),
        ("Chương 3", "Tấn công Rogue AP và Evil Twin là gì?"),
        ("Chương 4", "Gói tin Wi-Fi Beacon — Nguồn dữ liệu chính của dự án"),
        ("Chương 5", "25 Đặc trưng Wi-Fi sử dụng trong mô hình học máy"),
        ("Chương 6", "Học máy (Machine Learning) là gì?"),
        ("Chương 7", "5 Thuật toán học máy được sử dụng"),
        ("Chương 8", "Tập dữ liệu và quy trình tiền xử lý"),
        ("Chương 9", "Thực nghiệm và kết quả RQ1"),
        ("Chương 10", "Câu hỏi Nghiên cứu và Kết luận"),
    ]
    for code, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{code}: ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(title)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)

    page_break(doc)

    # ══════════════════════════════════════════════
    #  CHƯƠNG 1: MẠNG WI-FI
    # ══════════════════════════════════════════════
    add_h1(doc, "Chương 1: Mạng Wi-Fi là gì? Hoạt động ra sao?")

    add_h2(doc, "1.1  Wi-Fi là gì?")
    add_body(doc,
        "Wi-Fi (Wireless Fidelity) là công nghệ mạng không dây cho phép các thiết bị như máy tính, "
        "điện thoại, máy tính bảng kết nối Internet hoặc trao đổi dữ liệu với nhau mà không cần "
        "dây cáp vật lý. Wi-Fi hoạt động bằng cách truyền tín hiệu vô tuyến (sóng radio) giữa "
        "thiết bị đầu cuối và thiết bị phát sóng trung tâm gọi là Access Point (AP) hay Router."
    )
    add_note_box(doc, "Ví dụ quen thuộc: Khi bạn kết nối điện thoại vào mạng Wi-Fi tại nhà hoặc quán cà phê, điện thoại của bạn đang giao tiếp với Access Point (cái hộp phát Wi-Fi) qua sóng vô tuyến.")

    add_h2(doc, "1.2  Các chuẩn Wi-Fi phổ biến")
    add_body(doc, "Wi-Fi được chuẩn hóa bởi tổ chức IEEE (Institute of Electrical and Electronics Engineers) theo các thế hệ:")
    add_table(doc,
        ["Chuẩn", "Tên thương mại", "Tần số", "Tốc độ tối đa", "Phổ biến tại"],
        [
            ["802.11b", "Wi-Fi 1", "2.4 GHz", "11 Mbps",  "Thiết bị cũ trước 2005"],
            ["802.11g", "Wi-Fi 3", "2.4 GHz", "54 Mbps",  "Thiết bị đầu 2000s"],
            ["802.11n", "Wi-Fi 4", "2.4/5 GHz","600 Mbps","Phổ biến nhất hiện nay"],
            ["802.11ac","Wi-Fi 5", "5 GHz",   "3.5 Gbps", "Router hiện đại"],
            ["802.11ax","Wi-Fi 6", "2.4/5/6 GHz","9.6 Gbps","Router cao cấp mới"],
        ],
        col_widths=[Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.2), Inches(2.0)]
    )

    add_h2(doc, "1.3  Tần số và Kênh Wi-Fi")
    add_body(doc,
        "Wi-Fi hoạt động trên dải tần số radio nhất định. Ở băng tần 2.4 GHz có 13 kênh "
        "đánh số từ 1 đến 13, nhưng chỉ 3 kênh 1, 6, 11 là không chồng chéo nhau (non-overlapping). "
        "Khi hai AP phát sóng trên cùng một kênh ở gần nhau, chúng sẽ gây nhiễu cho nhau. "
        "Đây là một trong những đặc điểm mà bài dự án này khai thác để phát hiện AP giả mạo."
    )
    add_note_box(doc, "Hình dung: Băng tần Wi-Fi giống như các làn đường cao tốc — mỗi kênh là một làn đường. Hai xe (AP) chạy trên cùng một làn sẽ gây tai nạn (nhiễu tín hiệu).")

    page_break(doc)

    # ══════════════════════════════════════════════
    #  CHƯƠNG 2: ACCESS POINT
    # ══════════════════════════════════════════════
    add_h1(doc, "Chương 2: Access Point (AP) và các loại AP")

    add_h2(doc, "2.1  Access Point là gì?")
    add_body(doc,
        "Access Point (AP) là thiết bị phần cứng trung tâm của mạng Wi-Fi, có nhiệm vụ phát sóng "
        "tín hiệu không dây và đóng vai trò cầu nối giữa các thiết bị không dây và mạng có dây "
        "(Ethernet/Internet). Mỗi AP được nhận dạng duy nhất bằng hai thông tin chính:"
    )
    add_bullet(doc, "SSID (Service Set Identifier)", "Tên mạng Wi-Fi hiển thị cho người dùng, ví dụ 'Cafe_Free_Wifi'. Tên này dài tối đa 32 ký tự và có thể bị ẩn đi.")
    add_bullet(doc, "BSSID (Basic Service Set Identifier)", "Địa chỉ MAC (Media Access Control) vật lý 48-bit duy nhất của card mạng trong AP, định dạng HEX như 00:1A:2B:3C:4D:5E. Về lý thuyết địa chỉ này là duy nhất trên toàn cầu.")

    add_h2(doc, "2.2  Phân loại AP theo tính hợp lệ")
    add_table(doc,
        ["Loại AP", "Định nghĩa", "Mức độ nguy hiểm"],
        [
            ["AP hợp lệ (Legitimate AP)", "AP được quản trị viên hệ thống cài đặt, cấu hình và giám sát chính thức.", "Không nguy hiểm"],
            ["Rogue AP (AP giả mạo cơ bản)", "AP lạ kết nối vào mạng nội bộ mà không được phép của quản trị viên. Có thể do nhân viên tự cắm router cá nhân.", "⚠️ Trung bình"],
            ["Evil Twin AP (AP sao chép)", "AP được kẻ tấn công cài đặt để giả mạo hoàn toàn một AP hợp lệ — cùng tên SSID, có thể giả cả địa chỉ MAC.", "🚨 Rất nguy hiểm"],
            ["Soft AP (AP phần mềm)", "AP được tạo ra bằng phần mềm trên máy tính/Raspberry Pi (ví dụ: hostapd, airbase-ng). Thường dùng trong tấn công.", "🚨 Nguy hiểm"],
        ],
        col_widths=[Inches(1.5), Inches(3.0), Inches(1.8)]
    )
    doc.add_paragraph()

    page_break(doc)

    # ══════════════════════════════════════════════
    #  CHƯƠNG 3: TẤN CÔNG ROGUE AP & EVIL TWIN
    # ══════════════════════════════════════════════
    add_h1(doc, "Chương 3: Tấn công Rogue AP và Evil Twin là gì?")

    add_h2(doc, "3.1  Cơ chế tấn công Evil Twin chi tiết")
    add_body(doc,
        "Tấn công Evil Twin (AP Anh Em Song Sinh) là dạng tấn công mạng không dây nghiêm trọng "
        "nhất trong an ninh Wi-Fi hiện đại. Kẻ tấn công tạo ra một AP giả hoàn toàn giống AP thật "
        "về tên mạng (SSID) và thậm chí cả địa chỉ MAC (BSSID), khiến người dùng không thể phân biệt bằng mắt thường."
    )
    add_body(doc, "Các bước tấn công điển hình:")
    add_bullet(doc, "Bước 1 — Trinh sát", "Kẻ tấn công dùng công cụ quét Wi-Fi (Airodump-ng, Kismet) để thu thập SSID, BSSID, kênh của AP mục tiêu.")
    add_bullet(doc, "Bước 2 — Giải xác thực (Deauth Attack)", "Kẻ tấn công phát liên tục gói tin hủy kết nối (Deauthentication frames) giả mạo từ AP thật, buộc tất cả người dùng bị ngắt khỏi AP hợp lệ.")
    add_bullet(doc, "Bước 3 — Dựng AP giả", "Dùng công cụ phần mềm (hostapd, airbase-ng) hoặc thiết bị chuyên dụng (Wi-Fi Pineapple) tạo AP giả cùng SSID, cùng kênh hoặc kênh mạnh hơn.")
    add_bullet(doc, "Bước 4 — Chờ nạn nhân kết nối", "Thiết bị của nạn nhân tự động kết nối vào AP có tín hiệu mạnh nhất mà nó đã biết — và lần này lại là AP giả.")
    add_bullet(doc, "Bước 5 — Tấn công Man-in-the-Middle", "Toàn bộ lưu lượng của nạn nhân đi qua máy kẻ tấn công: đánh cắp mật khẩu, cookie đăng nhập, dữ liệu ngân hàng...")

    add_note_box(doc,
        "Thực tế nguy hiểm: Các tấn công Evil Twin hoàn toàn tự động với công cụ Wi-Fi Pineapple "
        "(giá ~$99) và thường xảy ra tại sân bay, quán cà phê, trung tâm thương mại — nơi người dùng "
        "hay kết nối Wi-Fi công cộng mà không kiểm tra kỹ.",
        color_hex="FFE0E0"
    )

    add_h2(doc, "3.2  Tại sao khó phát hiện bằng mắt thường?")
    add_body(doc,
        "Người dùng thông thường gần như không thể phát hiện Evil Twin vì:"
    )
    add_bullet(doc, "Tên mạng giống hệt", "SSID 'Airport_Free_WiFi' của kẻ tấn công trông giống 100% với AP thật.")
    add_bullet(doc, "Cường độ tín hiệu mạnh hơn", "Kẻ tấn công đặt thiết bị gần nạn nhân nên tín hiệu thậm chí còn mạnh hơn AP thật ở xa.")
    add_bullet(doc, "Kết nối tự động", "Điện thoại tự động kết nối vào mạng đã biết mà không hỏi người dùng.")
    add_bullet(doc, "Giao diện giống hệt", "Kẻ tấn công có thể dựng trang web giả giống trang đăng nhập của khách sạn/sân bay.")

    page_break(doc)

    # ══════════════════════════════════════════════
    #  CHƯƠNG 4: GÓI TIN BEACON
    # ══════════════════════════════════════════════
    add_h1(doc, "Chương 4: Gói tin Wi-Fi Beacon — Nguồn dữ liệu chính của dự án")

    add_h2(doc, "4.1  Gói tin Beacon là gì?")
    add_body(doc,
        "Gói tin Beacon (Beacon Frame) là gói tin quảng bá mà mỗi Access Point tự động gửi ra "
        "môi trường không khí mỗi 102.4 mili giây (gần 10 lần/giây) để thông báo sự tồn tại của mình. "
        "Đây là cơ chế 'nhịp tim' của AP — thiết bị xung quanh dựa vào Beacon để phát hiện các mạng Wi-Fi có sẵn."
    )
    add_body(doc,
        "Gói tin Beacon chứa rất nhiều thông tin kỹ thuật về AP phát sóng: tên mạng, địa chỉ MAC, "
        "kênh hoạt động, khả năng bảo mật, tốc độ hỗ trợ, loại phần cứng... Chính vì vậy, "
        "bằng cách phân tích gói tin Beacon, chúng ta có thể 'đọc được' rất nhiều đặc điểm của thiết bị phát."
    )
    add_note_box(doc,
        "Hình dung: Gói tin Beacon giống như tấm biển quảng cáo mà mỗi AP liên tục phát ra "
        "— trên tấm biển đó ghi đầy đủ thông tin về AP đó. Dự án này đọc các 'tấm biển' này để "
        "phân biệt AP thật và AP giả."
    )

    add_h2(doc, "4.2  Công cụ bắt gói tin — Wireshark và tcpdump")
    add_body(doc,
        "Để thu thập dữ liệu Beacon, chúng ta cần đặt card mạng Wi-Fi vào chế độ Monitor Mode "
        "(giám sát toàn bộ gói tin trong không khí) rồi dùng phần mềm phân tích gói tin:"
    )
    add_bullet(doc, "Wireshark", "Phần mềm miễn phí, giao diện đồ họa, có thể lọc và xem từng trường dữ liệu của mỗi gói tin.")
    add_bullet(doc, "tcpdump / tshark", "Công cụ dòng lệnh, dùng để thu thập dữ liệu hàng loạt và xuất ra file .csv để xử lý bằng Python.")

    add_h2(doc, "4.3  Cấu trúc của một Beacon Frame")
    add_body(doc, "Một gói tin Beacon có cấu trúc phân lớp như sau:")
    add_table(doc,
        ["Lớp", "Tên trường", "Ví dụ giá trị", "Ý nghĩa"],
        [
            ["RadioTap",  "RSSI (dbm_antsignal)", "-55 dBm",      "Cường độ tín hiệu khi gói tin đến card mạng"],
            ["RadioTap",  "Channel Frequency",   "2437 MHz",     "Tần số vật lý của kênh phát"],
            ["802.11 MAC","BSSID",               "00:1A:2B:...", "Địa chỉ MAC của AP phát"],
            ["802.11 MAC","SSID",                "My_Wifi",      "Tên mạng Wi-Fi"],
            ["802.11 MAC","Sequence Number",     "1024",         "Số thứ tự gói tin"],
            ["802.11 MAC","Beacon Interval",     "102",          "Tần suất phát Beacon (ms)"],
            ["Capability","Privacy bit",         "1",            "Mạng có mật khẩu bảo mật (1=có, 0=không)"],
            ["Capability","HT Support",          "1",            "Hỗ trợ chuẩn 802.11n tốc độ cao (1=có)"],
        ],
        col_widths=[Inches(0.9), Inches(1.5), Inches(1.2), Inches(2.5)]
    )

    page_break(doc)

    # ══════════════════════════════════════════════
    #  CHƯƠNG 5: 25 ĐẶC TRƯNG
    # ══════════════════════════════════════════════
    add_h1(doc, "Chương 5: 25 Đặc trưng Wi-Fi sử dụng trong mô hình học máy")

    add_h2(doc, "5.1  Đặc trưng là gì?")
    add_body(doc,
        "Trong học máy, đặc trưng (feature) là một biến số đầu vào mà mô hình dùng để học và dự đoán. "
        "Giống như khi bạn đoán một người là bác sĩ hay kỹ sư — bạn dựa vào các đặc trưng như: "
        "trang phục, dụng cụ mang theo, nơi làm việc... Tương tự, mô hình phát hiện Rogue AP "
        "dựa vào 25 đặc trưng từ gói tin Beacon để phân biệt AP thật và AP giả."
    )

    add_h2(doc, "5.2  Bảng tổng hợp 25 đặc trưng theo nhóm")
    add_body(doc, "Các đặc trưng được chia thành 4 nhóm chính theo tính chất vật lý/kỹ thuật:")

    # Group A
    add_h3(doc, "Nhóm A: Tín hiệu vật lý (Physical Signal)")
    add_table(doc,
        ["#", "Đặc trưng", "Giá trị mẫu (AP thật)", "Giá trị mẫu (AP giả)", "Ý nghĩa phân loại"],
        [
            ["1","RSSI",    "-52 dBm (ổn định)","-45 đến -85 dBm (nhảy lớn)","Sự ổn định tín hiệu theo thời gian"],
            ["2","Channel", "6 (cố định)",       "8 (bất thường)",             "Phát hiện kênh phát lạ"],
            ["25","rssi_std","0.8 dB (rất nhỏ)", "6.2 dB (rất lớn)",          "Độ lệch chuẩn RSSI — đo tính ổn định"],
        ],
        col_widths=[Inches(0.3), Inches(1.0), Inches(1.5), Inches(1.5), Inches(1.9)]
    )
    doc.add_paragraph()

    # Group B
    add_h3(doc, "Nhóm B: Cấu hình mạng & Bảo mật (Network Config & Security)")
    add_table(doc,
        ["#", "Đặc trưng", "Giá trị mẫu (AP thật)", "Giá trị mẫu (AP giả)", "Ý nghĩa phân loại"],
        [
            ["4","BeaconInterval","102 ms (chuẩn IEEE)", "100 ms hoặc 200 ms","Cấu hình phần mềm giả lập thường lệch"],
            ["5","Privacy",       "1 (có mật khẩu)",     "0 (mở hoàn toàn)",  "AP giả hay tạo mạng mở để bẫy"],
            ["15","IsHidden",     "0 (hiển thị công khai)","1 (ẩn tên mạng)", "AP giả hay ẩn để tránh quản trị"],
            ["20","UnusualChannel","0 (kênh thông thường)","1 (kênh 14...)",  "Kênh phát ngoài quy chuẩn"],
            ["21","Is_WEP",      "0 (dùng WPA2/WPA3)",  "1 (mã hóa WEP yếu)","Công cụ tấn công hay dùng WEP"],
            ["22","Is_Open",     "0 (có mật khẩu)",     "1 (không mật khẩu)","Mạng bẫy thường để mở"],
        ],
        col_widths=[Inches(0.3), Inches(1.2), Inches(1.5), Inches(1.5), Inches(1.7)]
    )
    doc.add_paragraph()

    # Group C
    add_h3(doc, "Nhóm C: Năng lực phần cứng (Hardware Capability)")
    add_table(doc,
        ["#", "Đặc trưng", "Giá trị mẫu (AP thật)", "Giá trị mẫu (AP giả)", "Ý nghĩa phân loại"],
        [
            ["8", "RateCount",    "8 (đầy đủ)",          "4 (nghèo nàn)",        "Card USB giả lập thiếu tốc độ"],
            ["9", "ExtRateCount", "4",                    "0",                    "Tốc độ mở rộng không hỗ trợ"],
            ["11","HasHT",        "1 (hỗ trợ 802.11n)",  "0 (không hỗ trợ)",     "Card Wi-Fi cũ không có HT"],
            ["12","HTChannelWidth","1 (40 MHz)",          "0 (20 MHz)",           "Băng thông anten hạn chế"],
            ["13","HTStreams",    "2 (MIMO 2x2)",          "1 (đơn anten)",        "Router thật hỗ trợ đa luồng MIMO"],
            ["14","HasExtCap",   "1 (tính năng nâng cao)","0 (bỏ trống)",         "Firmware của nhà sản xuất"],
            ["19","BSSID_LocalAdmin","0 (MAC nhà sx)",    "1 (MAC ngẫu nhiên phần mềm)","Đây là đặc trưng quan trọng nhất"],
        ],
        col_widths=[Inches(0.3), Inches(1.3), Inches(1.5), Inches(1.7), Inches(1.4)]
    )
    doc.add_paragraph()

    # Group D
    add_h3(doc, "Nhóm D: Số thứ tự & Cấu trúc gói tin (Sequence & Frame)")
    add_table(doc,
        ["#", "Đặc trưng", "Giá trị mẫu (AP thật)", "Giá trị mẫu (AP giả)", "Ý nghĩa phân loại"],
        [
            ["3", "SequenceNumber","Tăng đều: 100→101→102","Nhảy lộn: 100→15→102", "Phát hiện chèn gói tin giả"],
            ["16","FrameLength",  "128 bytes (cố định)",  "142 bytes (thêm tag lạ)","Công cụ tấn công thêm payload"],
            ["17","SSID_Length",  "14 ký tự",             "0 ký tự (trống)",        "Gói tin giả hay để tên trống"],
            ["18","EmptySSID",    "0 (có tên)",            "1 (SSID trống)",          "Gói tin bẫy không tên mạng"],
            ["23","LowSeqNumber", "0 (sequence bình thường)","1 (< 10, khởi động lại)","Công cụ tấn công thường bị restart"],
            ["24","seq_delta",   "±1 (tăng đều)",         "±2500 (nhảy rất lớn)",   "Dấu hiệu Evil Twin chạy song song"],
        ],
        col_widths=[Inches(0.3), Inches(1.2), Inches(1.5), Inches(1.7), Inches(1.5)]
    )

    page_break(doc)

    # ══════════════════════════════════════════════
    #  CHƯƠNG 6: HỌC MÁY
    # ══════════════════════════════════════════════
    add_h1(doc, "Chương 6: Học máy (Machine Learning) là gì?")

    add_h2(doc, "6.1  Định nghĩa đơn giản")
    add_body(doc,
        "Học máy (Machine Learning - ML) là một nhánh của Trí tuệ nhân tạo (AI), "
        "nơi chúng ta huấn luyện máy tính học từ dữ liệu thay vì lập trình các quy tắc thủ công. "
        "Thay vì viết code 'NẾU privacy=0 VÀ IsHidden=1 VÀ... THÌ là rogue', "
        "chúng ta cho máy tính tự học mối liên hệ phức tạp giữa hàng chục đặc trưng và nhãn phân loại."
    )
    add_note_box(doc,
        "Ví dụ dễ hiểu: Học máy giống như dạy một đứa trẻ phân biệt chó và mèo — "
        "bạn không cần giải thích công thức, chỉ cần cho xem hàng ngàn bức ảnh chó và mèo đã được gán nhãn. "
        "Từ đó đứa trẻ (mô hình) tự rút ra đặc điểm phân biệt."
    )

    add_h2(doc, "6.2  Học có giám sát (Supervised Learning)")
    add_body(doc,
        "Dự án này sử dụng phương pháp Học có giám sát — loại học máy mà dữ liệu huấn luyện "
        "đã có sẵn nhãn đúng/sai (ground truth label). Mỗi dòng dữ liệu có:"
    )
    add_bullet(doc, "Đặc trưng đầu vào (X)", "25 cột đặc trưng Wi-Fi từ gói tin Beacon.")
    add_bullet(doc, "Nhãn đầu ra (y)", "Cột is_rogue: 0 = AP hợp lệ, 1 = Rogue AP.")
    add_body(doc, "Quy trình học có giám sát trong dự án:")
    add_bullet(doc, "Bước 1", "Chia dữ liệu: 80% để huấn luyện (train) / 20% để kiểm tra (test).")
    add_bullet(doc, "Bước 2", "Mô hình học từ 80% dữ liệu đã có nhãn, tìm ra quy luật phân loại.")
    add_bullet(doc, "Bước 3", "Dùng mô hình đã học để dự đoán nhãn của 20% dữ liệu còn lại.")
    add_bullet(doc, "Bước 4", "So sánh dự đoán với nhãn thật để tính độ chính xác.")

    add_h2(doc, "6.3  Các chỉ số đánh giá mô hình")
    add_body(doc, "Để đo hiệu năng mô hình phân loại, chúng ta dùng ma trận nhầm lẫn (Confusion Matrix):")
    add_table(doc,
        ["", "Dự đoán: Thật (0)", "Dự đoán: Rogue (1)"],
        [
            ["Thực tế: Thật (0)", "True Negative (TN)\nDự đoán đúng là AP thật", "False Positive (FP)\nBáo nhầm AP thật là rogue"],
            ["Thực tế: Rogue (1)", "False Negative (FN)\nBỏ sót Rogue AP nguy hiểm", "True Positive (TP)\nPhát hiện đúng Rogue AP"],
        ],
        col_widths=[Inches(1.8), Inches(2.2), Inches(2.2)]
    )
    doc.add_paragraph()
    add_body(doc, "Từ ma trận nhầm lẫn, tính ra 4 chỉ số đánh giá quan trọng:")
    add_table(doc,
        ["Chỉ số", "Công thức", "Ý nghĩa trong bài toán này"],
        [
            ["Accuracy (Độ chính xác)", "TP+TN / Tổng", "Tỷ lệ dự đoán đúng tổng thể"],
            ["Precision (Độ chính xác dương)", "TP / (TP+FP)", "Trong số báo rogue, bao nhiêu % thực sự là rogue?"],
            ["Recall (Độ nhạy)", "TP / (TP+FN)", "Trong số Rogue thật, bao nhiêu % được phát hiện?"],
            ["F1-Score", "2×P×R / (P+R)", "Trung bình hài hòa của Precision và Recall — quan trọng nhất khi dữ liệu mất cân bằng"],
            ["AUC-ROC", "Diện tích dưới đường cong ROC", "Khả năng phân biệt tổng thể: 0.5=ngẫu nhiên, 1.0=hoàn hảo"],
        ],
        col_widths=[Inches(1.7), Inches(1.5), Inches(3.0)]
    )
    add_note_box(doc,
        "Trong bài toán bảo mật, Recall quan trọng hơn Precision: bỏ sót 1 Rogue AP (FN) "
        "nguy hiểm hơn báo nhầm 1 AP thật (FP). Vì vậy chúng ta chú trọng F1-Score và Recall."
    )

    page_break(doc)

    # ══════════════════════════════════════════════
    #  CHƯƠNG 7: 5 THUẬT TOÁN
    # ══════════════════════════════════════════════
    add_h1(doc, "Chương 7: 5 Thuật toán Học máy được sử dụng")

    # KNN
    add_h2(doc, "7.1  K-Nearest Neighbors (KNN) — Thuật toán láng giềng gần nhất")
    add_body(doc,
        "KNN là thuật toán đơn giản nhất: để phân loại một điểm dữ liệu mới, "
        "nó tìm K điểm dữ liệu gần nhất trong tập train (K=5 trong dự án này) "
        "và bỏ phiếu đa số — nếu 4/5 láng giềng là Rogue thì điểm mới cũng được phân loại là Rogue."
    )
    add_note_box(doc,
        "Ví dụ thực tế: Bạn muốn biết một quán ăn ngon hay không — hãy hỏi 5 người bạn sống gần đó nhất. "
        "Nếu 4/5 người bảo ngon → quán đó ngon. KNN hoạt động y hệt vậy."
    )
    add_bullet(doc, "Ưu điểm", "Không cần giả thuyết phân phối dữ liệu, trực quan, dễ hiểu.")
    add_bullet(doc, "Nhược điểm", "Rất chậm khi dự đoán (phải tính khoảng cách tới mọi điểm train) — inference time 2.33s trong thực nghiệm. Kém khi dữ liệu nhiều chiều.")
    add_bullet(doc, "Kết quả trong dự án", "Accuracy 88.98%, F1-Score 82.61%.")

    # SVM
    add_h2(doc, "7.2  Support Vector Machine (SVM) — Máy vector hỗ trợ")
    add_body(doc,
        "SVM tìm một siêu phẳng (hyperplane) tối ưu để phân chia ranh giới giữa hai lớp "
        "với biên giới rộng nhất (Maximum Margin Hyperplane). Dự án dùng LinearSVC — "
        "phiên bản SVM tuyến tính, phù hợp với dữ liệu lớn."
    )
    add_note_box(doc,
        "Ví dụ: Hãy tưởng tượng các điểm AP thật và AP giả được vẽ lên mặt phẳng 2D. "
        "SVM vẽ một đường thẳng phân chia hai nhóm sao cho khoảng cách từ đường thẳng "
        "đến điểm gần nhất của mỗi nhóm là lớn nhất."
    )
    add_bullet(doc, "Ưu điểm", "Hiệu quả tốt với dữ liệu tuyến tính, ít tham số cần điều chỉnh.")
    add_bullet(doc, "Nhược điểm", "Kém khi ranh giới phân lớp phi tuyến (thực tế với Wi-Fi). Recall chỉ 63.84% — bỏ sót nhiều Rogue AP.")
    add_bullet(doc, "Kết quả trong dự án", "Accuracy 84.48%, F1-Score 74.00% — thấp nhất trong 5 mô hình.")

    # Random Forest
    add_h2(doc, "7.3  Random Forest — Rừng cây ngẫu nhiên")
    add_body(doc,
        "Random Forest xây dựng nhiều cây quyết định (Decision Tree) độc lập, mỗi cây "
        "được huấn luyện trên một tập con ngẫu nhiên của dữ liệu và đặc trưng. Kết quả cuối cùng "
        "là kết quả bỏ phiếu đa số của 100 cây (n_estimators=100 trong dự án)."
    )
    add_note_box(doc,
        "Ví dụ: Thay vì hỏi ý kiến 1 chuyên gia (1 cây quyết định), bạn hỏi ý kiến 100 chuyên gia "
        "độc lập rồi chọn ý kiến đa số. Kết quả ổn định và chính xác hơn nhiều so với 1 chuyên gia."
    )
    add_bullet(doc, "Ưu điểm", "Chống Overfitting tốt, xử lý tốt dữ liệu phi tuyến, tự động tính Feature Importance.")
    add_bullet(doc, "Nhược điểm", "Chậm hơn các mô hình Boosting, kém interpretable hơn một cây đơn lẻ.")
    add_bullet(doc, "Kết quả trong dự án", "Accuracy 93.47%, F1-Score 90.03%, AUC-ROC 0.9743.")

    # XGBoost
    add_h2(doc, "7.4  XGBoost — Extreme Gradient Boosting (Mô hình tốt nhất)")
    add_body(doc,
        "XGBoost là thuật toán Boosting: xây dựng từng cây quyết định tuần tự, "
        "mỗi cây mới học từ sai lầm của cây trước. Khác với Random Forest (các cây độc lập song song), "
        "XGBoost cải thiện liên tục theo từng vòng (iteration). Đây là thuật toán phổ biến nhất "
        "trong các cuộc thi Kaggle và nghiên cứu học máy thực tế."
    )
    add_note_box(doc,
        "Hình dung: Random Forest = 100 học sinh làm bài độc lập, lấy điểm trung bình. "
        "XGBoost = 1 học sinh làm bài, thầy chấm và chỉ ra lỗi, học sinh sửa, thầy chấm lại... "
        "lặp 100 lần. Kết quả của XGBoost chính xác hơn nhờ quá trình học lặp liên tục."
    )
    add_bullet(doc, "Ưu điểm", "Độ chính xác cao nhất, xử lý tốt giá trị thiếu, rất nhanh khi dự đoán (0.02s), hỗ trợ tính Feature Importance và SHAP.")
    add_bullet(doc, "Nhược điểm", "Nhiều siêu tham số cần điều chỉnh (hyperparameter tuning).")
    add_bullet(doc, "Kết quả trong dự án", "Accuracy 94.10%, F1-Score 90.99%, AUC-ROC 0.9754 — tốt nhất trong 5 mô hình.")

    # LightGBM
    add_h2(doc, "7.5  LightGBM — Light Gradient Boosting Machine")
    add_body(doc,
        "LightGBM do Microsoft phát triển, là phiên bản tối ưu hóa tốc độ của XGBoost. "
        "Thay vì chia cây theo chiều ngang (level-wise như XGBoost), LightGBM chia theo chiều lá "
        "(leaf-wise) giúp hội tụ nhanh hơn với dữ liệu lớn."
    )
    add_bullet(doc, "Ưu điểm", "Nhanh nhất trong các thuật toán Boosting, tiêu thụ ít RAM hơn XGBoost, phù hợp dataset lớn triệu dòng.")
    add_bullet(doc, "Nhược điểm", "Có thể Overfit hơn XGBoost với dữ liệu nhỏ nếu không điều chỉnh tham số.")
    add_bullet(doc, "Kết quả trong dự án", "Accuracy 93.69%, F1-Score 90.27%, AUC-ROC 0.9733 — tương đương XGBoost, Precision cao nhất (96.86%).")

    page_break(doc)

    # ══════════════════════════════════════════════
    #  CHƯƠNG 8: TẬP DỮ LIỆU
    # ══════════════════════════════════════════════
    add_h1(doc, "Chương 8: Tập dữ liệu và Quy trình tiền xử lý")

    add_h2(doc, "8.1  Mô tả tập dữ liệu clean_dataset.csv")
    add_body(doc,
        "Tập dữ liệu sử dụng trong thực nghiệm là clean_dataset.csv — được xây dựng từ việc "
        "thu thập gói tin Beacon thực tế trong môi trường phòng thí nghiệm mạng không dây. "
        "Dữ liệu bao gồm cả gói tin từ các AP hợp lệ thật và các AP giả mạo được tạo ra có chủ đích."
    )
    add_table(doc,
        ["Thông số", "Giá trị"],
        [
            ["Tổng số dòng (mẫu)", "67.776 dòng gói tin Beacon"],
            ["Số cột (đặc trưng)", "31 cột (25 đặc trưng + cột nhãn is_rogue + một số cột phụ)"],
            ["Số mẫu AP hợp lệ (is_rogue=0)", "~44.000 mẫu (65%)"],
            ["Số mẫu Rogue AP (is_rogue=1)", "~23.776 mẫu (35%)"],
            ["Tỷ lệ giá trị thiếu (NaN)", "Có — chủ yếu ở các đặc trưng HT khi AP không hỗ trợ 802.11n"],
            ["Định dạng", "CSV, mã hóa UTF-8"],
        ],
        col_widths=[Inches(2.5), Inches(3.7)]
    )

    add_h2(doc, "8.2  Quy trình tiền xử lý dữ liệu")
    add_body(doc, "Trước khi đưa vào huấn luyện mô hình, dữ liệu cần qua các bước xử lý sau:")
    add_bullet(doc, "Bước 1 — Loại bỏ rò rỉ thời gian (Leakage Prevention)",
        "Xác định và loại bỏ 5 cột chứa thông tin thời gian thu thập dữ liệu: "
        "Timestamp_ms, BeaconTimestamp, LowTSF, Timestamp_Ratio, time_delta. "
        "Nếu giữ lại các cột này, mô hình đạt ~100% accuracy ảo vì nó học thời điểm chụp "
        "gói tin chứ không học đặc điểm của AP!")
    add_bullet(doc, "Bước 2 — Xử lý giá trị thiếu (Missing Value Handling)",
        "Loại bỏ các dòng có giá trị NaN bằng dropna(). "
        "Sau bước này còn lại 45.424 dòng sạch hoàn toàn.")
    add_bullet(doc, "Bước 3 — Chuẩn hóa dữ liệu (Feature Scaling)",
        "Chỉ áp dụng cho KNN và SVM (nhạy cảm với thang đo): dùng StandardScaler "
        "chuyển đổi mỗi đặc trưng về trung bình 0 và độ lệch chuẩn 1. "
        "Quan trọng: Scaler chỉ được FIT trên tập Train để tránh rò rỉ thông tin test vào train.")
    add_bullet(doc, "Bước 4 — Phân chia Train/Test (80/20)",
        "Phân chia ngẫu nhiên có phân tầng (stratified): đảm bảo tỷ lệ Rogue/Legitimate "
        "trong Train set và Test set giống nhau với tập gốc.")

    add_h2(doc, "8.3  Lý giải tại sao phải loại bỏ cột Timestamp")
    add_body(doc,
        "Đây là vấn đề kỹ thuật quan trọng nhất trong thực nghiệm. Khi thu thập dữ liệu trong lab, "
        "tất cả Rogue AP thường được bật lên trong một phiên làm việc cụ thể "
        "(ví dụ: 14:00-15:00 ngày 01/06/2025). Mô hình học được rằng 'timestamp trong khoảng đó = rogue' "
        "thay vì học các đặc trưng vật lý thực sự. Kết quả: accuracy 100% trong lab nhưng 0% trong thực tế!"
    )
    add_note_box(doc,
        "Thực tế: Đây gọi là Data Leakage (Rò rỉ dữ liệu) — một trong những lỗi nghiêm trọng nhất "
        "trong học máy. Giống như cho học sinh xem đáp án trước khi thi — điểm thi rất cao "
        "nhưng thực tế không học được gì!",
        color_hex="FFE0E0"
    )

    page_break(doc)

    # ══════════════════════════════════════════════
    #  CHƯƠNG 9: KẾT QUẢ
    # ══════════════════════════════════════════════
    add_h1(doc, "Chương 9: Thực nghiệm và Kết quả RQ1")

    add_h2(doc, "9.1  Thiết lập thực nghiệm")
    add_table(doc,
        ["Thông số", "Giá trị"],
        [
            ["Tập dữ liệu", "clean_dataset.csv (45.424 dòng sau dropna)"],
            ["Số đặc trưng đầu vào", "25 đặc trưng (loại bỏ 5 cột rò rỉ)"],
            ["Tỷ lệ phân chia", "Train 80% = 36.339 dòng / Test 20% = 9.085 dòng"],
            ["Phương pháp phân chia", "Stratified Random Split (random_state=42)"],
            ["Môi trường chạy", "Python 3.10, scikit-learn 1.3, xgboost 2.0, lightgbm 4.0"],
            ["Phần cứng", "CPU Intel Core i7, RAM 16GB"],
        ],
        col_widths=[Inches(2.5), Inches(3.7)]
    )

    add_h2(doc, "9.2  Kết quả so sánh 5 mô hình (Objective 1.2)")
    add_table(doc,
        ["Thuật toán", "Accuracy", "Precision", "Recall", "F1-Score", "AUC-ROC", "Train time", "Predict time"],
        [
            ["KNN",          "88.98%","91.01%","75.64%","82.61%","0.9204","0.04s","2.33s"],
            ["SVM (Linear)", "84.48%","88.03%","63.84%","74.00%","0.8844","0.44s","0.00s"],
            ["Random Forest","93.47%","95.54%","85.11%","90.03%","0.9743","0.49s","0.12s"],
            ["XGBoost ⭐",   "94.10%","96.54%","86.04%","90.99%","0.9754","0.37s","0.02s"],
            ["LightGBM",     "93.69%","96.86%","84.51%","90.27%","0.9733","0.48s","0.03s"],
        ],
        col_widths=[Inches(1.3), Inches(0.8), Inches(0.8), Inches(0.7), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8)]
    )
    doc.add_paragraph()

    add_h2(doc, "9.3  Phân tích kết quả (Objective 1.1 — Feature Importance)")
    add_body(doc, "Top 5 đặc trưng quan trọng nhất theo phân tích XGBoost và SHAP:")
    add_table(doc,
        ["Hạng", "Đặc trưng", "Độ quan trọng", "Lý giải"],
        [
            ["1", "IsHidden",       "28.5%", "AP giả hay ẩn SSID để tránh bị quản trị viên phát hiện."],
            ["2", "Privacy",        "16.0%", "AP giả tạo mạng mở (không mật khẩu) để dụ nạn nhân."],
            ["3", "Is_WEP",         "9.7%",  "Công cụ tấn công dùng mã hóa WEP cũ làm mặc định."],
            ["4", "BSSID_LocalAdmin","6.9%", "Địa chỉ MAC ngẫu nhiên phần mềm → cờ LocalAdmin bật."],
            ["5", "HTStreams",       "6.7%", "Card tấn công chỉ có 1 luồng MIMO, Router thật có 2-4."],
        ],
        col_widths=[Inches(0.5), Inches(1.5), Inches(1.0), Inches(3.2)]
    )

    # Embed images
    img_model = "plots/model_comparison.png"
    img_feat = "plots/feature_importance.png"
    img_shap = "plots/shap_summary.png"

    doc.add_paragraph()
    if os.path.exists(img_model):
        doc.add_picture(img_model, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph("Hình 1: So sánh hiệu năng 5 thuật toán học máy (không rò rỉ dữ liệu)")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True
        cap.runs[0].font.size = Pt(10)

    doc.add_paragraph()
    if os.path.exists(img_feat):
        doc.add_picture(img_feat, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph("Hình 2: Độ quan trọng của các đặc trưng Wi-Fi trong mô hình XGBoost")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True
        cap.runs[0].font.size = Pt(10)

    doc.add_paragraph()
    if os.path.exists(img_shap):
        doc.add_picture(img_shap, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph("Hình 3: Biểu đồ SHAP phân tích cơ chế ra quyết định của mô hình")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True
        cap.runs[0].font.size = Pt(10)

    page_break(doc)

    # ══════════════════════════════════════════════
    #  CHƯƠNG 10: CÂU HỎI NGHIÊN CỨU & KẾT LUẬN
    # ══════════════════════════════════════════════
    add_h1(doc, "Chương 10: Câu hỏi Nghiên cứu và Kết luận")

    add_h2(doc, "10.1  Câu hỏi Nghiên cứu 1 (RQ1) — Trả lời")
    add_body(doc,
        "RQ1: Làm thế nào các mô hình học máy phân biệt chính xác AP hợp lệ và AP giả mạo/evil twin "
        "bằng các đặc trưng lưu lượng và cường độ tín hiệu Wi-Fi?"
    )
    add_body(doc, "Trả lời:")
    add_bullet(doc, "Objective 1.1 — Đặc trưng hiệu quả nhất",
        "Nhóm đặc trưng cấu hình mạng bảo mật (IsHidden, Privacy, Is_WEP) và đặc trưng "
        "phần cứng (BSSID_LocalAdmin, HTStreams) là những chỉ số quan trọng nhất. "
        "Chúng phản ánh sự khác biệt cơ bản giữa phần cứng Router thương mại thật "
        "và thiết bị/phần mềm giả lập của kẻ tấn công.")
    add_bullet(doc, "Objective 1.2 — Thuật toán tốt nhất",
        "XGBoost là thuật toán hiệu quả nhất với Accuracy 94.10%, F1-Score 90.99% "
        "và thời gian dự đoán chỉ 0.02 giây — khả thi để triển khai thời gian thực.")

    add_h2(doc, "10.2  Tại sao XGBoost tốt nhất?")
    add_body(doc, "XGBoost vượt trội vì:")
    add_bullet(doc, "Dữ liệu phi tuyến", "Ranh giới giữa AP thật và giả phức tạp, phi tuyến — cây quyết định xử lý tốt hơn mô hình tuyến tính (SVM).")
    add_bullet(doc, "Boosting liên tục", "Mỗi vòng học cải thiện sai lầm vòng trước — chính xác hơn Random Forest (song song độc lập).")
    add_bullet(doc, "Xử lý NaN tốt", "XGBoost tự nội tại xử lý giá trị thiếu, không cần bước điền khuyết phức tạp.")

    add_h2(doc, "10.3  Hạn chế và hướng phát triển")
    add_bullet(doc, "Hạn chế", "Tập dữ liệu thu thập trong môi trường lab có thể không phản ánh đầy đủ sự đa dạng của mạng Wi-Fi thực tế.")
    add_bullet(doc, "Hướng 1", "Thử nghiệm trên dataset_2 (35 file CSV, ~1.75 triệu dòng) — đa dạng hơn nhưng cần pipeline tiền xử lý phức tạp hơn.")
    add_bullet(doc, "Hướng 2", "Thêm 5-Fold Cross-Validation để có kết quả trung bình ổn định với khoảng tin cậy (±).")
    add_bullet(doc, "Hướng 3", "Triển khai mô hình thời gian thực trên Raspberry Pi kết hợp với công cụ quét Wi-Fi.")

    add_h2(doc, "10.4  Tổng kết kiến thức cần nắm")
    add_table(doc,
        ["Chủ đề", "Điểm mấu chốt cần nhớ"],
        [
            ["Wi-Fi Beacon", "Mỗi AP phát 10 gói Beacon/giây chứa đầy thông tin kỹ thuật — đây là nguồn dữ liệu của dự án."],
            ["Evil Twin", "AP giả sao chép SSID/BSSID để đánh lừa người dùng — nguy hiểm vì không thể phân biệt bằng mắt."],
            ["25 đặc trưng", "Chia 4 nhóm: Tín hiệu vật lý, Cấu hình bảo mật, Năng lực phần cứng, Số thứ tự gói tin."],
            ["Data Leakage", "Phải loại bỏ cột Timestamp — không loại bỏ sẽ cho kết quả 100% ảo!"],
            ["F1-Score", "Chỉ số quan trọng nhất khi dữ liệu mất cân bằng và bài toán bảo mật ưu tiên Recall."],
            ["XGBoost tốt nhất", "94.10% Accuracy, 0.02s dự đoán, học từ sai lầm liên tục — phù hợp triển khai thực tế."],
        ],
        col_widths=[Inches(2.0), Inches(4.2)]
    )

    doc.add_paragraph()
    add_note_box(doc,
        "Kết luận cuối: Với chỉ 25 đặc trưng từ gói tin Beacon sẵn có trong mọi mạng Wi-Fi, "
        "mô hình XGBoost có thể phát hiện Rogue AP với độ chính xác 94.10% mà không cần "
        "phần cứng đặc biệt hay sự can thiệp thủ công của quản trị viên — "
        "mở ra hướng xây dựng hệ thống phát hiện tự động giá thành thấp.",
        color_hex="E8F5E9"
    )

    # ══════════════════════════════════════════════
    #  LƯU
    # ══════════════════════════════════════════════
    out_path = "Theory_RogueAP_Concepts.docx"
    doc.save(out_path)
    print(f"[OK] Tai lieu ly thuyet da duoc luu: {out_path}")

if __name__ == "__main__":
    build()
